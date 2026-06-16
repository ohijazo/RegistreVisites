#!/usr/bin/env python3
"""Genera el manual del recepcionista en format .docx.

Ús:
    pip install python-docx
    python scripts/build_manual.py

Sortida: docs/manual/Manual_Recepcionista_Visites.docx

Les captures de pantalla es busquen a docs/manual/img/. Si no existeixen,
s'insereix un placeholder en gris i el document es pot regenerar a mesura
que es van afegint imatges.
"""
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent
IMG_DIR = ROOT / "docs" / "manual" / "img"
OUTPUT = ROOT / "docs" / "manual" / "Manual_Recepcionista_Visites.docx"

BASE_URL = "http://visitesfc.agrienergia.local"
ADMIN_URL = f"{BASE_URL}/admin/"

GREY = RGBColor(0x88, 0x88, 0x88)
BLUE = RGBColor(0x1F, 0x4E, 0x79)


# ---------- helpers ----------

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = BLUE
    return p


def h2(doc, text):
    return doc.add_heading(text, level=2)


def h3(doc, text):
    return doc.add_heading(text, level=3)


def para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def numbered(doc, text):
    return doc.add_paragraph(text, style="List Number")


def code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    return p


def url(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = BLUE
    return p


def screenshot(doc, filename, caption=""):
    path = IMG_DIR / filename
    if path.exists():
        try:
            doc.add_picture(str(path), width=Inches(6.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as exc:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f"[ERROR carregant {filename}: {exc}]")
            r.italic = True
            r.font.color.rgb = GREY
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"[CAPTURA PENDENT: {filename}]")
        r.italic = True
        r.font.color.rgb = GREY
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = GREY


def table_two_col(doc, header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=2)
    t.style = "Light Grid Accent 1"
    hc = t.rows[0].cells
    for i, txt in enumerate(header):
        hc[i].text = ""
        run = hc[i].paragraphs[0].add_run(txt)
        run.bold = True
    for r_idx, row in enumerate(rows, start=1):
        cells = t.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
    return t


def table_perms(doc):
    """Taula de permisos del rol receptionist vs admin/viewer."""
    rows = [
        ("Veure dashboard de visites actives", "Sí", "Sí", "Sí"),
        ("Cercar a l'historial", "Sí", "Sí", "Sí"),
        ("Veure detall d'una visita", "Sí", "Sí", "Sí"),
        ("Veure DNI desxifrat (cal contrasenya)", "Sí", "Sí", "No"),
        ("Registrar sortida manual", "Sí", "Sí", "No"),
        ("Exportar a Excel/CSV", "Sí", "Sí", "No"),
        ("Crear i gestionar visites previstes", "Sí", "Sí", "No"),
        ("Veure estadístiques", "Sí", "Sí", "Sí"),
        ("Editar hores d'entrada/sortida", "Sí", "No", "No"),
        ("Eliminar una visita (dret RGPD)", "Sí", "No", "No"),
        ("Gestionar departaments / textos legals / usuaris", "Sí", "No", "No"),
        ("Veure logs d'auditoria", "Sí", "No", "No"),
    ]
    t = doc.add_table(rows=1 + len(rows), cols=4)
    t.style = "Light Grid Accent 1"
    hc = t.rows[0].cells
    for i, txt in enumerate(("Funcionalitat", "Admin", "Recepcionista", "Viewer")):
        hc[i].text = ""
        run = hc[i].paragraphs[0].add_run(txt)
        run.bold = True
    for r_idx, row in enumerate(rows, start=1):
        cells = t.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
    return t


def page_break(doc):
    doc.add_page_break()


# ---------- contingut ----------

def build():
    doc = Document()

    # Estils per defecte
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # ============== PORTADA ==============
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Registre de Visites")
    r.bold = True
    r.font.size = Pt(32)
    r.font.color.rgb = BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Manual del recepcionista")
    r.font.size = Pt(20)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Accés al panell d'administració:")
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(ADMIN_URL)
    r.bold = True
    r.font.name = "Consolas"
    r.font.size = Pt(14)
    r.font.color.rgb = BLUE

    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Versió {date.today():%Y-%m-%d}")
    r.italic = True
    r.font.color.rgb = GREY

    page_break(doc)

    # ============== 1. INTRODUCCIÓ ==============
    h1(doc, "1. Introducció")

    h2(doc, "1.1 Què és aquesta aplicació")
    para(doc,
         "Aplicació web interna que substitueix el llibre de registre de visites en paper. "
         "Els visitants s'identifiquen a una tablet de recepció (o escanejant un QR amb el "
         "mòbil), accepten les normes d'accés i el tractament de les seves dades, i queden "
         "registrats al sistema. Quan marxen, registren la sortida introduint el seu DNI a "
         "la mateixa tablet."
         )
    para(doc,
         "El teu paper com a recepcionista és supervisar aquest registre des del navegador, "
         "ajudar els visitants si tenen dubtes, registrar sortides quan algú marxi sense "
         "fer-ho a la tablet, i consultar o exportar les dades que necessitin altres "
         "departaments."
         )

    h2(doc, "1.2 El teu rol: què pots fer i què no")
    para(doc,
         "Tens el rol \"receptionist\". A continuació es resumeix què et permet aquest rol, "
         "comparat amb un administrador i amb un usuari de només lectura:"
         )
    table_perms(doc)
    para(doc, "")
    para(doc,
         "Si necessites fer alguna acció marcada com a \"No\" (eliminar una visita, modificar "
         "hores d'entrada/sortida, gestionar usuaris), demana-ho a l'administrador del "
         "sistema."
         )

    h2(doc, "1.3 Com accedir-hi")
    para(doc, "Obre el navegador i ves a:")
    url(doc, ADMIN_URL)
    para(doc,
         "Et demanaran l'email i la contrasenya que l'administrador t'haurà facilitat. "
         "Si no recordes la contrasenya, contacta amb l'administrador perquè te la "
         "reestableixi: no hi ha funció d'autoreset per email."
         )

    page_break(doc)

    # ============== 2. PRIMER ACCÉS ==============
    h1(doc, "2. Primer accés")

    h2(doc, "2.1 Login")
    para(doc,
         "A la pantalla de login, introdueix el teu email i la contrasenya. Si les "
         "credencials són correctes, accediràs al panell principal (Dashboard)."
         )
    screenshot(doc, "01_login.png", "Pantalla de login del panell d'administració")
    para(doc, "Punts a tenir en compte:")
    bullet(doc, "L'email és el que t'ha donat l'administrador, normalment el corporatiu.")
    bullet(doc, "La sessió dura unes hores; si la deixes oberta tot el dia probablement "
                "hauràs de tornar a entrar al matí següent.")
    bullet(doc, "Després de diversos intents fallits seguits, el sistema bloqueja "
                "temporalment el login per seguretat. Espera uns minuts i torna-ho a "
                "provar.")

    h2(doc, "2.2 Navegació general")
    para(doc,
         "Un cop dins, tens una barra de navegació amb les diferents seccions. Les opcions "
         "marcades amb (admin) només surten si tens el rol d'administrador i no apareixen "
         "per a recepcionistes."
         )
    screenshot(doc, "02_layout.png", "Barra de navegació i estructura general del panell")
    para(doc, "Seccions principals que veuràs:")
    bullet(doc, "Dashboard — visites actives en temps real i resum del dia.")
    bullet(doc, "Historial — totes les visites passades, amb filtres i exportació.")
    bullet(doc, "Visites previstes — gestió de visites planificades amb antelació.")
    bullet(doc, "Estadístiques — gràfics i resums per al període que triïs.")

    h2(doc, "2.3 Tancar sessió i canviar contrasenya")
    para(doc,
         "Per tancar sessió, clica al teu nom a la part superior dreta i tria \"Tancar "
         "sessió\". Des del mateix menú també pots accedir al teu perfil per canviar la "
         "contrasenya."
         )
    screenshot(doc, "03_profile.png", "Menú d'usuari i perfil")
    para(doc, "Si canvies la contrasenya:")
    bullet(doc, "Han de ser mínim 12 caràcters.")
    bullet(doc, "Utilitza una contrasenya única, no la reutilitzis d'altres serveis.")
    bullet(doc, "Si la oblides, l'administrador la pot resetejar; no hi ha email d'autoservei.")

    page_break(doc)

    # ============== 3. DASHBOARD ==============
    h1(doc, "3. Dashboard — la pantalla del dia a dia")
    para(doc,
         "El Dashboard és la pantalla principal que utilitzaràs constantment. Mostra qui "
         "hi ha actualment a les instal·lacions, quines visites estan previstes per avui "
         "i un resum del moviment del dia. S'actualitza automàticament cada pocs segons."
         )

    h2(doc, "3.1 Visió general")
    screenshot(doc, "10_dashboard.png", "Vista completa del Dashboard")

    h2(doc, "3.2 Visites previstes (bàner superior)")
    para(doc,
         "Si algú de l'oficina ha registrat una visita prevista per avui, apareix al "
         "bàner superior. Així saps qui s'espera i pots avisar el departament corresponent "
         "quan el visitant arribi. Quan el visitant es registra al quiosc, la previsió es "
         "marca automàticament com a arribada."
         )
    screenshot(doc, "11_dashboard_expected.png",
               "Bàner de visites previstes pendents d'arribar")

    h2(doc, "3.3 Targetes de resum")
    para(doc, "A la part superior tens quatre indicadors del dia:")
    bullet(doc, "Visites actives ara — gent dins de les instal·lacions en aquest moment.")
    bullet(doc, "Entrades avui — total de registres d'entrada des de mitjanit.")
    bullet(doc, "Sortides avui — total de sortides registrades.")
    bullet(doc, "Durada mitjana — temps mitjà que han durat les visites completes d'avui.")

    h2(doc, "3.4 Taula de visites actives + colors d'alerta")
    para(doc,
         "Al cos del dashboard hi ha la taula amb tots els visitants que encara no han "
         "registrat la sortida, ordenats per hora d'entrada. La columna \"Temps dins\" "
         "indica quant fa que han entrat."
         )
    screenshot(doc, "12_active_table.png", "Taula de visites actives amb colors d'alerta")
    para(doc, "Codi de colors:")
    bullet(doc, "Fila normal: visita en curs sense incidència.")
    bullet(doc, "Fila en groc: visita que dura més del normal (a partir de poques hores).")
    bullet(doc, "Fila en vermell: visita anormalment llarga; probablement la persona ha "
                "marxat sense registrar la sortida.")
    para(doc,
         "Si veus visites en vermell, fes una sortida manual (secció següent) o avisa "
         "l'administrador si no estàs segur de què ha passat."
         )

    h2(doc, "3.5 Registrar una sortida manual des del dashboard")
    para(doc,
         "Cada fila de la taula té un botó per registrar la sortida sense necessitat que "
         "el visitant torni a la tablet. Útil quan algú marxa per una porta diferent o "
         "se n'ha oblidat. La fila desapareix immediatament després."
         )
    screenshot(doc, "13_manual_checkout.png", "Botó de sortida manual a la taula d'actives")

    page_break(doc)

    # ============== 4. HISTORIAL ==============
    h1(doc, "4. Historial de visites")
    para(doc, "URL: ")
    url(doc, f"{BASE_URL}/admin/visits")

    h2(doc, "4.1 Vista general i columnes")
    para(doc,
         "L'historial mostra totes les visites registrades, ordenades per data d'entrada "
         "(les més recents a dalt). Per defecte es mostren les més recents; si vols veure "
         "visites antigues, utilitza els filtres."
         )
    screenshot(doc, "20_visits_list.png", "Llistat de visites a l'historial")
    para(doc, "Columnes principals:")
    bullet(doc, "Nom i cognoms — clica per veure el detall complet.")
    bullet(doc, "Empresa.")
    bullet(doc, "Departament visitat.")
    bullet(doc, "Hora d'entrada i de sortida.")
    bullet(doc, "Durada de la visita (només si està completada).")
    bullet(doc, "Estat: \"Activa\" si encara no ha sortit, \"Completada\" si sí.")

    h2(doc, "4.2 Filtres")
    para(doc, "Pots combinar diversos filtres a la part superior:")
    bullet(doc, "Rang de dates (des de / fins a).")
    bullet(doc, "Empresa — text lliure, busca aproximat.")
    bullet(doc, "Departament — desplegable amb els departaments configurats.")
    bullet(doc, "Nom o cognoms — text lliure.")
    bullet(doc, "Estat — Totes / Actives / Completades.")
    screenshot(doc, "21_filters.png", "Filtres de cerca a l'historial")
    para(doc,
         "Clica \"Cercar\" per aplicar els filtres i \"Netejar\" per descartar-los i "
         "tornar al llistat per defecte."
         )

    h2(doc, "4.3 Paginació i ordenació")
    para(doc,
         "Els resultats es paginen en blocs de 25. Pots navegar amb els botons de pàgina "
         "i ordenar per qualsevol columna clicant a la capçalera."
         )

    h2(doc, "4.4 Veure el detall d'una visita")
    para(doc,
         "Clica al nom del visitant (o al botó \"Veure\") per accedir al detall complet, "
         "explicat a la secció 5."
         )

    page_break(doc)

    # ============== 5. DETALL D'UNA VISITA ==============
    h1(doc, "5. Detall d'una visita")
    para(doc, "URL d'exemple: ")
    url(doc, f"{BASE_URL}/admin/visits/<id>")

    h2(doc, "5.1 Dades del visitant")
    para(doc,
         "Mostra el nom, cognoms, empresa, telèfon (si l'ha indicat) i idioma en què va "
         "fer el registre."
         )
    screenshot(doc, "30_visit_detail.png", "Vista del detall d'una visita")

    h2(doc, "5.2 Dades de la visita i acceptació RGPD")
    para(doc,
         "També veus quin departament visitava, el motiu indicat, hora exacta d'entrada i "
         "de sortida, durada total, i la versió del document legal que va acceptar amb el "
         "timestamp i la IP. Si va deixar signatura, també es mostra."
         )

    h2(doc, "5.3 Veure el DNI (només quan calgui)")
    para(doc,
         "El DNI es guarda xifrat i no es mostra per defecte. Si necessites consultar-lo "
         "(per a un tema legal, una incidència o petició d'autoritat), clica \"Veure DNI\"."
         )
    screenshot(doc, "31_view_dni.png", "Botó i diàleg per desxifrar el DNI")
    para(doc, "Important:", bold=True)
    bullet(doc, "Et demanarà la teva contrasenya per confirmar.")
    bullet(doc, "El DNI es mostra pocs segons i s'amaga sola.")
    bullet(doc, "Cada vegada que ho fas, queda un registre a l'auditoria amb el teu nom, "
                "data i hora. No s'utilitza com a control; serveix per traçabilitat si "
                "algun dia s'investiga un accés concret.")

    h2(doc, "5.4 Registrar sortida manual des del detall")
    para(doc,
         "Si la visita encara està activa, des del detall també pots registrar la sortida "
         "manualment, igual que des del dashboard."
         )

    h2(doc, "5.5 Imprimir o desar com a PDF")
    para(doc,
         "El navegador et permet imprimir el detall o desar-lo com a PDF (Ctrl+P). Útil "
         "quan algú demana una còpia del registre d'una visita concreta."
         )

    page_break(doc)

    # ============== 6. EXPORTAR ==============
    h1(doc, "6. Exportar dades")

    h2(doc, "6.1 Exportar a Excel o CSV")
    para(doc,
         "Des del llistat de l'historial pots descarregar el resultat com a fitxer Excel "
         "o CSV. Els botons són a la barra superior, al costat dels filtres."
         )
    screenshot(doc, "40_export.png", "Botons d'exportació Excel/CSV a l'historial")

    h2(doc, "6.2 Els filtres apliquen a l'exportació")
    para(doc,
         "L'exportació respecta exactament els filtres que tens aplicats. Si filtres per "
         "\"empresa = Otis\" i \"data des de = 01/01\", l'Excel només contindrà aquests "
         "registres. Si no apliques cap filtre, exportarà el llistat actual (amb un límit "
         "màxim per protegir el rendiment del servidor)."
         )

    h2(doc, "6.3 Què s'inclou i què no")
    para(doc, "L'Excel/CSV inclou:")
    bullet(doc, "Nom i cognoms, empresa, telèfon.")
    bullet(doc, "Departament i motiu de la visita.")
    bullet(doc, "Idioma, data i hora d'entrada i sortida, durada.")
    bullet(doc, "Mètode de sortida (QR / DNI / manual) i timestamp d'acceptació RGPD.")
    para(doc, "El que NO s'exporta mai:", bold=True)
    bullet(doc, "El DNI — ni xifrat ni en clar. Si necessites el DNI d'un visitant concret, "
                "consulta'l individualment al detall (secció 5.3).")

    page_break(doc)

    # ============== 7. VISITES PREVISTES ==============
    h1(doc, "7. Visites previstes")
    para(doc, "URL: ")
    url(doc, f"{BASE_URL}/admin/expected")

    h2(doc, "7.1 Per a què serveixen")
    para(doc,
         "Si algú de l'oficina sap que demà vindrà un client, un proveïdor o un consultor, "
         "pot crear una \"visita prevista\". Això té dos avantatges:"
         )
    bullet(doc, "Apareix al teu Dashboard al bàner superior, així saps qui esperar.")
    bullet(doc, "Es genera un codi d'accés que es pot enviar al visitant per email. Quan "
                "el visitant arriba i introdueix el codi al quiosc, no ha d'omplir tot el "
                "formulari: ja venen pre-omplertes les dades.")
    screenshot(doc, "50_expected_list.png", "Llistat de visites previstes")

    h2(doc, "7.2 Crear una nova visita prevista")
    para(doc,
         "Clica \"+ Nova visita prevista\" al llistat. Cal indicar el nom del visitant, "
         "l'empresa, el departament que visitarà, el motiu i la data/hora aproximada. "
         "També pots indicar l'amfitrió (qui rep la visita)."
         )
    screenshot(doc, "51_expected_new.png", "Formulari de nova visita prevista")
    para(doc, "Camps clau a destacar:", bold=True)
    bullet(doc, "Email del visitant — si el poses, després podràs enviar-li la invitació "
                "amb el codi de pre-registre directament des del sistema (vegeu 7.6).")
    bullet(doc, "Amfitrió — apareix a la notificació interna i facilita filtrar al "
                "dashboard \"les meves visites\" per a aquesta persona.")
    bullet(doc, "Estat inicial \"Pendent\" — un cop el visitant arribi al quiosc i es "
                "registri, passa automàticament a \"Arribada\".")

    h2(doc, "7.3 Vista calendari")
    para(doc,
         "A més del llistat tradicional, pots veure les visites previstes en un calendari "
         "mensual o setmanal, útil per planificar la cobertura de recepció els dies de "
         "molt moviment."
         )
    screenshot(doc, "52_expected_calendar.png", "Vista de calendari de visites previstes")

    h2(doc, "7.4 Marcar com arribada o cancel·lar")
    para(doc,
         "Si el visitant arriba i es registra al quiosc, la previsió es marca "
         "automàticament com a arribada. Si avisa que no vindrà o canvia el dia, pots "
         "marcar-la com a cancel·lada des del detall de la previsió."
         )

    h2(doc, "7.5 Compartir el codi d'accés manualment")
    para(doc,
         "Cada visita prevista té un codi de 8 caràcters i un QR associat. Pots "
         "compartir-lo manualment (copia el codi amb el botó \"Copiar\", o descarrega el "
         "QR amb \"Descarregar QR\") i enviar-lo per WhatsApp, missatge corporatiu o "
         "qualsevol altre canal. Al quiosc, l'opció \"Tinc un codi de visita\" permet "
         "introduir-lo i el formulari surt pre-omplert."
         )

    h2(doc, "7.6 Enviar notificacions per email des del sistema")
    para(doc,
         "Des del detall d'una visita prevista (estat \"Pendent\") tens dos botons "
         "d'enviament d'email, amb propòsits diferents:"
         )
    screenshot(doc, "53_expected_notify.png",
               "Detall d'una visita prevista amb els botons d'enviament d'email")

    h3(doc, "Botó morat \"Enviar invitació al visitant\"")
    para(doc,
         "Envia un email al visitant amb el codi d'accés i el QR per fer pre-registre "
         "des del seu mòbil abans d'arribar. Requereix tenir omplert el camp \"Email del "
         "visitant\" al detall — si no, en lloc del botó veuràs un missatge indicant que "
         "l'has d'afegir clicant \"Editar\"."
         )
    para(doc, "Si l'email ja s'ha enviat un cop, el botó canvia a \"Reenviar invitació "
              "al visitant\" i sota es mostra quan es va enviar i a qui."
         )

    h3(doc, "Botó blau \"Enviar notificació\"")
    para(doc,
         "Envia un email intern als amfitrions o persones de l'oficina indicades, amb "
         "les dades de la previsió (visitant, empresa, dia, hora, motiu). És el mateix "
         "email que es genera automàticament quan es crea la visita prevista; aquest botó "
         "serveix per reenviar-lo o enviar-lo a destinataris addicionals."
         )
    para(doc,
         "En clicar el botó s'obre un diàleg on pots editar la llista de destinataris "
         "(separats per comes) abans d'enviar. L'assumpte i el cos del missatge no són "
         "editables; són els que el sistema genera per defecte."
         )

    h3(doc, "Si els botons surten desactivats")
    para(doc,
         "Si veus els botons en gris i no pots clicar-los, és perquè el servidor encara "
         "no té configurat el backend d'enviament d'email. Avisa l'administrador del "
         "sistema perquè ho activi al fitxer de configuració."
         )

    page_break(doc)

    # ============== 8. ESTADÍSTIQUES ==============
    h1(doc, "8. Estadístiques")
    para(doc, "URL: ")
    url(doc, f"{BASE_URL}/admin/stats")

    h2(doc, "8.1 Resum del període")
    para(doc,
         "Tria un rang de dates a la part superior i tindràs un resum: total de visites, "
         "visitants únics, empreses úniques, durada mitjana i visites que no van registrar "
         "sortida."
         )
    screenshot(doc, "60_stats_top.png", "Resum d'estadístiques per al període")

    h2(doc, "8.2 Gràfic de visites per dia")
    para(doc,
         "Un gràfic de barres mostra quantes visites hi va haver cada dia del període. "
         "Permet detectar pics i tendències."
         )
    screenshot(doc, "61_stats_daily.png", "Gràfic de visites per dia")

    h2(doc, "8.3 Gràfic per departament i per franja horària")
    para(doc,
         "També tens un gràfic circular amb la distribució per departament i un gràfic de "
         "barres per franja horària. El segon és útil per planificar quan necessites més "
         "cobertura a recepció."
         )

    h2(doc, "8.4 Top empreses")
    para(doc,
         "Una taula resumeix quines empreses han visitat més vegades al període i quan va "
         "ser l'última visita."
         )

    page_break(doc)

    # ============== 9. QUÈ VEU EL VISITANT ==============
    h1(doc, "9. Què veu el visitant")
    para(doc,
         "Aquesta secció explica el flux del visitant al quiosc de recepció (tablet) o al "
         "seu mòbil. Et serveix per ajudar-lo si té dubtes o problemes."
         )

    h2(doc, "9.1 Selecció d'idioma")
    para(doc,
         "Quan el visitant s'apropa a la tablet, veu primer una pantalla per triar idioma "
         "entre català, castellà, francès i anglès."
         )
    screenshot(doc, "70_visitor_language.png", "Pantalla inicial de selecció d'idioma")

    h2(doc, "9.2 Menú: Entrar, Sortir o Codi")
    para(doc, "A continuació pot triar entre tres opcions:")
    bullet(doc, "Registrar entrada — quan acaba d'arribar.")
    bullet(doc, "Registrar sortida — quan marxa.")
    bullet(doc, "Tinc un codi de visita — si ha rebut un codi de pre-registre per email.")
    screenshot(doc, "71_visitor_action.png", "Menú amb les tres opcions principals")

    h2(doc, "9.3 Formulari de dades")
    para(doc,
         "Si tria \"Registrar entrada\", se li demana: nom, cognoms, empresa, DNI/NIE/"
         "passaport, departament que visita, motiu i telèfon (opcional). Si ja havia "
         "visitat abans, en escriure el DNI el sistema reomple automàticament les dades."
         )
    screenshot(doc, "72_visitor_form.png", "Formulari de dades del visitant")

    h2(doc, "9.4 Acceptació de normes i signatura")
    para(doc,
         "El visitant ha de llegir el text de normes i RGPD (es fa scroll obligatori fins "
         "al final), marcar les acceptacions corresponents i signar amb el dit a la "
         "pantalla. Sense signatura no pot completar el registre."
         )
    screenshot(doc, "73_visitor_legal.png", "Acceptació de normes i signatura")

    h2(doc, "9.5 Confirmació d'entrada")
    para(doc,
         "Un cop completat, veu una confirmació amb el seu nom i l'hora d'entrada. La "
         "pantalla es reinicia sola després d'uns segons perquè el següent visitant la "
         "trobi neta."
         )
    screenshot(doc, "74_visitor_confirmation.png", "Pantalla de confirmació d'entrada")

    h2(doc, "9.6 Registre de sortida")
    para(doc,
         "Quan marxa, el visitant tria \"Registrar sortida\" al menú i introdueix el seu "
         "DNI. El sistema reconeix la visita activa i registra la sortida automàticament."
         )
    screenshot(doc, "75_visitor_checkout.png", "Pantalla de registre de sortida amb DNI")

    h2(doc, "9.7 Si ha rebut un codi de pre-registre")
    para(doc,
         "Si l'oficina li havia enviat un codi (vegeu secció 7), pot triar \"Tinc un codi "
         "de visita\", introduir-lo, i el formulari surt pre-omplert. Així estalvia temps "
         "i evita errors d'escriptura del DNI o del nom de l'empresa."
         )
    screenshot(doc, "76_visitor_code.png", "Pantalla d'introducció del codi de pre-registre")

    page_break(doc)

    # ============== 10. PREGUNTES FREQÜENTS ==============
    h1(doc, "10. Preguntes freqüents")

    h3(doc, "Un visitant ha marxat sense registrar sortida")
    para(doc,
         "Ves al detall de la visita (Dashboard o Historial) i clica \"Sortida manual\". "
         "Quedarà registrat amb mètode \"manual\" perquè quedi traça que no va sortir pel "
         "procediment habitual."
         )

    h3(doc, "No trobo una visita antiga")
    para(doc,
         "Comprova els filtres de data a l'historial. Per defecte només es mostren les més "
         "recents. Amplia el rang \"des de\" cap enrere."
         )

    h3(doc, "Necessito el DNI d'un visitant per a una incidència o autoritat")
    para(doc,
         "Entra al detall de la visita i clica \"Veure DNI\". Et demanarà la contrasenya. "
         "L'acció queda enregistrada a l'auditoria."
         )

    h3(doc, "Vull avisar que demà vindrà X")
    para(doc,
         "Crea una visita prevista (secció 7). Si vols que el visitant pugui registrar-se "
         "ràpid, comparteix-li el codi de 8 caràcters per email."
         )

    h3(doc, "He fet una sortida manual per error")
    para(doc,
         "Els recepcionistes no podem editar hores d'entrada/sortida. Avisa l'administrador "
         "del sistema perquè ho corregeixi."
         )

    h3(doc, "Vull les visites del mes per a un informe")
    para(doc,
         "A l'historial, filtra pel rang de dates del mes i clica \"Excel\". El fitxer es "
         "descarregarà amb totes les columnes excepte el DNI."
         )

    h3(doc, "L'aplicació no respon o dóna error")
    para(doc,
         "Avisa l'equip d'IT. Inclou una captura de pantalla si és possible. L'aplicació "
         "s'executa al servidor intern de l'empresa; les incidències es resolen des d'allà."
         )

    page_break(doc)

    # ============== 11. CONTACTE ==============
    h1(doc, "11. Contacte")
    para(doc,
         "Per a qualsevol problema amb l'aplicació, accés bloquejat, reset de contrasenya "
         "o petició que no puguis fer amb el teu rol, contacta amb l'administrador del "
         "sistema."
         )
    para(doc, "URL del panell d'administració:")
    url(doc, ADMIN_URL)

    # ============== Desa ==============
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


def main():
    out = build()
    expected = [
        "01_login.png", "02_layout.png", "03_profile.png",
        "10_dashboard.png", "11_dashboard_expected.png",
        "12_active_table.png", "13_manual_checkout.png",
        "20_visits_list.png", "21_filters.png",
        "30_visit_detail.png", "31_view_dni.png",
        "40_export.png",
        "50_expected_list.png", "51_expected_new.png", "52_expected_calendar.png",
        "53_expected_notify.png",
        "60_stats_top.png", "61_stats_daily.png",
        "70_visitor_language.png", "71_visitor_action.png", "72_visitor_form.png",
        "73_visitor_legal.png", "74_visitor_confirmation.png",
        "75_visitor_checkout.png", "76_visitor_code.png",
    ]
    missing = [f for f in expected if not (IMG_DIR / f).exists()]
    print(f"Manual generat: {out}")
    print(f"Captures total: {len(expected)} | trobades: {len(expected) - len(missing)} | pendents: {len(missing)}")
    if missing:
        print("Captures pendents:")
        for f in missing:
            print(f"  - {f}")


if __name__ == "__main__":
    main()
